#!/usr/bin/env python3

import sys
import os
from pathlib import Path

# Add etl_pipeline to path for imports
sys.path.append('etl_pipeline')

def create_test_xlsx():
    """Create a comprehensive test XLSX file"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        from datetime import datetime, date
        
        print("Creating comprehensive test XLSX file...")
        
        # Create workbook
        wb = Workbook()
        wb.remove(wb.active)
        
        # Sheet 1: Employee Data with formulas
        ws1 = wb.create_sheet(title="Employee_Data")
        
        # Headers
        headers = ['ID', 'Name', 'Department', 'Hire_Date', 'Base_Salary', 'Bonus', 'Total_Compensation', 'Performance']
        for col, header in enumerate(headers, 1):
            cell = ws1.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        
        # Employee data with formulas
        employees = [
            (1001, "Alice Johnson", "Engineering", date(2022, 3, 15), 75000, 7500),
            (1002, "Bob Smith", "Marketing", date(2021, 7, 22), 62000, 6200),
            (1003, "Carol Davis", "Engineering", date(2023, 1, 10), 78000, 8500),
            (1004, "David Wilson", "Sales", date(2020, 11, 5), 55000, 8000),
            (1005, "Emma Brown", "HR", date(2022, 9, 18), 58000, 4500)
        ]
        
        performance_scores = [4.2, 3.8, 4.5, 4.1, 3.9]
        
        for row, (emp_data, perf) in enumerate(zip(employees, performance_scores), 2):
            for col, value in enumerate(emp_data, 1):
                ws1.cell(row=row, column=col, value=value)
            
            # Add total compensation formula
            ws1.cell(row=row, column=7, value=f"=E{row}+F{row}")
            ws1.cell(row=row, column=8, value=perf)
        
        # Summary formulas
        ws1.cell(row=8, column=1, value="SUMMARY:")
        ws1.cell(row=8, column=2, value="Total Employees")
        ws1.cell(row=8, column=3, value="=COUNTA(A2:A6)")
        
        ws1.cell(row=9, column=2, value="Avg Base Salary")
        ws1.cell(row=9, column=3, value="=AVERAGE(E2:E6)")
        
        ws1.cell(row=10, column=2, value="Total Compensation")
        ws1.cell(row=10, column=3, value="=SUM(G2:G6)")
        
        # Sheet 2: Sales Data with Charts
        ws2 = wb.create_sheet(title="Sales_Data")
        
        quarters = ['Q1 2025', 'Q2 2025', 'Q3 2025', 'Q4 2025', 'Total']
        products = ['Product A', 'Product B', 'Product C', 'Product D', 'TOTAL']
        
        # Headers
        for col, quarter in enumerate(quarters, 2):
            ws2.cell(row=1, column=col, value=quarter)
        
        # Product rows with data
        sales_data = [
            [125000, 132000, 128000, 145000],
            [98000, 105000, 112000, 118000],
            [87000, 91000, 95000, 102000],
            [156000, 168000, 172000, 189000]
        ]
        
        for row, (product, sales) in enumerate(zip(products[:-1], sales_data), 2):
            ws2.cell(row=row, column=1, value=product)
            for col, value in enumerate(sales, 2):
                ws2.cell(row=row, column=col, value=value)
            # Total formula
            ws2.cell(row=row, column=6, value=f"=SUM(B{row}:E{row})")
        
        # Totals row
        ws2.cell(row=6, column=1, value="TOTAL")
        for col in range(2, 7):
            ws2.cell(row=6, column=col, value=f"=SUM({chr(64+col)}2:{chr(64+col)}5)")
        
        # Sheet 3: Analytics with complex formulas
        ws3 = wb.create_sheet(title="Analytics")
        
        ws3.cell(row=1, column=1, value="ETL Pipeline Analytics Dashboard")
        ws3.cell(row=1, column=1).font = Font(size=16, bold=True)
        
        # Cross-sheet references
        ws3.cell(row=3, column=1, value="Employee Metrics:")
        ws3.cell(row=4, column=1, value="Total Employees:")
        ws3.cell(row=4, column=2, value="=Employee_Data!C8")
        
        ws3.cell(row=5, column=1, value="Average Salary:")
        ws3.cell(row=5, column=2, value="=Employee_Data!C9")
        
        ws3.cell(row=6, column=1, value="Total Payroll:")
        ws3.cell(row=6, column=2, value="=Employee_Data!C10")
        
        ws3.cell(row=8, column=1, value="Sales Metrics:")
        ws3.cell(row=9, column=1, value="Q1 Total Sales:")
        ws3.cell(row=9, column=2, value="=Sales_Data!B6")
        
        ws3.cell(row=10, column=1, value="Annual Sales:")
        ws3.cell(row=10, column=2, value="=Sales_Data!F6")
        
        ws3.cell(row=12, column=1, value="Performance Analysis:")
        ws3.cell(row=13, column=1, value="High Performers (>4.0):")
        ws3.cell(row=13, column=2, value='=COUNTIF(Employee_Data!H2:H6,">4")')
        
        ws3.cell(row=14, column=1, value="Average Performance:")
        ws3.cell(row=14, column=2, value="=AVERAGE(Employee_Data!H2:H6)")
        
        # Date functions
        ws3.cell(row=16, column=1, value="Report Generated:")
        ws3.cell(row=16, column=2, value="=NOW()")
        
        # Sheet 4: Data Validation and Formatting
        ws4 = wb.create_sheet(title="Formatted_Data")
        
        ws4.cell(row=1, column=1, value="Formatted Data Test")
        ws4.cell(row=1, column=1).font = Font(size=14, bold=True, color="FF0000")
        
        # Different number formats
        ws4.cell(row=3, column=1, value="Currency:")
        ws4.cell(row=3, column=2, value=1234.56)
        ws4.cell(row=3, column=2).number_format = '"$"#,##0.00'
        
        ws4.cell(row=4, column=1, value="Percentage:")
        ws4.cell(row=4, column=2, value=0.85)
        ws4.cell(row=4, column=2).number_format = '0.0%'
        
        ws4.cell(row=5, column=1, value="Date:")
        ws4.cell(row=5, column=2, value=date.today())
        
        ws4.cell(row=6, column=1, value="Time:")
        ws4.cell(row=6, column=2, value=datetime.now())
        
        # Conditional formatting example
        for i in range(8, 13):
            ws4.cell(row=i, column=1, value=f"Score {i-7}:")
            score = (i-7) * 0.8 + 1.5
            cell = ws4.cell(row=i, column=2, value=score)
            
            if score >= 4.0:
                cell.fill = PatternFill(start_color="00FF00", end_color="00FF00", fill_type="solid")
            elif score >= 3.0:
                cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
            else:
                cell.fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")
        
        # Save the workbook
        test_cases_dir = Path('test_cases')
        xlsx_path = test_cases_dir / 'test_spreadsheet.xlsx'
        wb.save(str(xlsx_path))
        
        print(f"✅ XLSX file created: {xlsx_path}")
        return True
        
    except ImportError:
        print("❌ openpyxl not available - skipping XLSX creation")
        return False
    except Exception as e:
        print(f"❌ Error creating XLSX: {e}")
        return False

def create_test_docx():
    """Create a comprehensive test DOCX file"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        print("Creating comprehensive test DOCX file...")
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('ETL Pipeline Comprehensive Test Document', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Introduction
        intro = doc.add_paragraph()
        intro_run1 = intro.add_run('This document serves as a comprehensive test case for the ETL pipeline\'s ')
        intro_run2 = intro.add_run('enhanced DOCX handler')
        intro_run2.bold = True
        intro_run3 = intro.add_run('. It contains various document elements to validate structure preservation, ')
        intro_run4 = intro.add_run('content extraction')
        intro_run4.italic = True
        intro_run5 = intro.add_run(', and metadata generation capabilities.')
        
        # Section 1: Document Structure
        doc.add_heading('1. Document Structure Testing', level=2)
        
        doc.add_paragraph(
            'This section validates the handler\'s ability to preserve hierarchical document structure. '
            'The enhanced DOCX handler should correctly identify heading levels, maintain content organization, '
            'and generate appropriate metadata for each section.'
        )
        
        doc.add_heading('1.1 Nested Heading Structure', level=3)
        doc.add_paragraph('This subsection tests nested heading recognition and hierarchical content organization.')
        
        doc.add_heading('1.2 Content Types', level=3)
        doc.add_paragraph('The following content types should be properly identified and processed:')
        
        # Lists
        doc.add_paragraph('Text Processing Features:', style='List Bullet')
        doc.add_paragraph('Structure preservation with heading hierarchy', style='List Bullet')
        doc.add_paragraph('Rich content extraction including formatting', style='List Bullet')
        doc.add_paragraph('Table data extraction and organization', style='List Bullet')
        doc.add_paragraph('List item identification and grouping', style='List Bullet')
        
        # Section 2: Content Formatting
        doc.add_heading('2. Content Formatting and Styles', level=2)
        
        # Content Formatting
        para = doc.add_paragraph('This paragraph demonstrates various ')
        bold_run = para.add_run('text formatting')
        bold_run.bold = True
        para.add_run(' options including ')
        bold_run2 = para.add_run('bold text')
        bold_run2.bold = True
        para.add_run(', ')
        italic_run = para.add_run('italic text')
        italic_run.italic = True
        para.add_run(', and ')
        under_run = para.add_run('underlined content')
        under_run.underline = True
        para.add_run('. The handler should extract content while noting formatting metadata.')
        
        # Numbered list
        doc.add_paragraph('Processing Steps:', style='List Number')
        doc.add_paragraph('Load document and parse structure', style='List Number')
        doc.add_paragraph('Extract content from each element', style='List Number')
        doc.add_paragraph('Generate metadata for content types', style='List Number')
        doc.add_paragraph('Create structured output with regions', style='List Number')
        doc.add_paragraph('Validate extraction accuracy and completeness', style='List Number')
        
        # Section 3: Table Processing
        doc.add_heading('3. Table Data Extraction', level=2)
        
        doc.add_paragraph('The following table tests the handler\'s ability to extract tabular data:')
        
        # Create comprehensive table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Component', 'Type', 'Status', 'Performance', 'Notes']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # Data rows
        test_data = [
            ('Text Handler', 'Core', 'Stable', 'Excellent', 'Multi-format support'),
            ('Binary Handler', 'Enhanced', 'Active', 'Very Good', 'XLSX/DOCX processing'),
            ('Mixed Handler', 'Optimized', 'Active', 'Good', 'Content type detection'),
            ('OCR Integration', 'External', 'Available', 'Good', 'Chandra OCR service'),
            ('Format Detection', 'Core', 'Optimized', 'Excellent', 'MIME type classification'),
            ('Metadata Extraction', 'Core', 'Enhanced', 'Very Good', 'Rich metadata generation')
        ]
        
        for component, type_, status, performance, notes in test_data:
            row_cells = table.add_row().cells
            row_cells[0].text = component
            row_cells[1].text = type_
            row_cells[2].text = status
            row_cells[3].text = performance
            row_cells[4].text = notes
        
        # Section 4: Complex Content
        doc.add_heading('4. Complex Content Processing', level=2)
        
        doc.add_heading('4.1 Mixed Content Elements', level=3)
        
        quote_para = doc.add_paragraph()
        quote_para.add_run('"The ETL pipeline demonstrates exceptional capability in processing ')
        quote_italic = quote_para.add_run('diverse document formats')
        quote_italic.italic = True
        quote_para.add_run(' while maintaining ')
        quote_bold = quote_para.add_run('data integrity')
        quote_bold.bold = True
        quote_para.add_run(' and structure preservation."')
        
        doc.add_heading('4.2 Performance Metrics', level=3)
        
        # Performance table
        perf_table = doc.add_table(rows=1, cols=3)
        perf_table.style = 'Light List Accent 1'
        
        perf_headers = ['Metric', 'Value', 'Target']
        for i, header in enumerate(perf_headers):
            perf_table.rows[0].cells[i].text = header
        
        perf_data = [
            ('Processing Speed', '< 500ms per document', '< 1000ms'),
            ('Accuracy Rate', '98.5%', '> 95%'),
            ('Structure Preservation', '97.2%', '> 90%'),
            ('Memory Usage', '45MB average', '< 100MB')
        ]
        
        for metric, value, target in perf_data:
            row_cells = perf_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
            row_cells[2].text = target
        
        # Section 5: Conclusion
        doc.add_heading('5. Processing Validation', level=2)
        
        conclusion = doc.add_paragraph(
            'This comprehensive test document validates multiple aspects of the enhanced DOCX handler. '
            'Expected processing results include structured content regions, preserved heading hierarchy, '
            'extracted table data, and rich metadata for each content type. The handler should demonstrate '
            'robust processing capabilities while maintaining document structure integrity.'
        )
        
        final_para = doc.add_paragraph()
        final_bold = final_para.add_run('Processing Status: ')
        final_bold.bold = True
        final_para.add_run('This document contains ')
        final_italic1 = final_para.add_run('15+ content sections')
        final_italic1.italic = True
        final_para.add_run(' with ')
        final_italic2 = final_para.add_run('2 tables')
        final_italic2.italic = True
        final_para.add_run(', ')
        final_italic3 = final_para.add_run('3 heading levels')
        final_italic3.italic = True
        final_para.add_run(', and ')
        final_italic4 = final_para.add_run('multiple formatting styles')
        final_italic4.italic = True
        final_para.add_run(' for comprehensive validation.')
        
        # Save the document
        test_cases_dir = Path('test_cases')
        docx_path = test_cases_dir / 'test_document.docx'
        doc.save(str(docx_path))
        
        print(f"✅ DOCX file created: {docx_path}")
        return True
        
    except ImportError:
        print("❌ python-docx not available - skipping DOCX creation")
        return False
    except Exception as e:
        print(f"❌ Error creating DOCX: {e}")
        return False

if __name__ == "__main__":
    print("=== Creating Comprehensive Test Files ===")
    
    docx_success = create_test_docx()
    xlsx_success = create_test_xlsx()
    
    print(f"\n=== Test File Creation Summary ===")
    if docx_success:
        print("✅ DOCX test file created with comprehensive content")
    if xlsx_success:
        print("✅ XLSX test file created with multiple worksheets and formulas")
    
    print(f"\n📁 All test files available in: test_cases/")