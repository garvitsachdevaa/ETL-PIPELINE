import logging
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
import io
import uuid
from typing import List, Dict, Any, Optional
from ingestion.schemas import DocumentObject
from handlers.binary_schema import BinaryDocument, Page, Block, Region

logger = logging.getLogger(__name__)

def extract_docx_text(raw_bytes: bytes) -> Dict[str, Any]:
    """Extract comprehensive content from DOCX file bytes with structure preservation"""
    try:
        doc = Document(io.BytesIO(raw_bytes))
        
        # Extract structured document elements
        document_structure = []
        
        # Process all document elements in order
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para_data = _extract_paragraph_structure(element, doc)
                if para_data:
                    document_structure.append(para_data)
            elif element.tag.endswith('tbl'):  # Table
                table_data = _extract_table_structure(element)
                if table_data:
                    document_structure.append(table_data)
        
        # Also extract using python-docx objects for backup
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                para_info = _get_paragraph_info(para)
                paragraphs.append(para_info)
        
        # Extract tables using python-docx
        tables = []
        for table in doc.tables:
            table_data = _extract_table_data(table)
            if table_data:
                tables.append(table_data)
        
        # Generate structured text that preserves hierarchy
        structured_text = _generate_structured_text(document_structure)
        
        # Fallback to simple text if structured extraction fails
        if not structured_text.strip():
            all_text = []
            for para in paragraphs:
                all_text.append(para['text'])
            for table in tables:
                table_text = _table_to_text(table)
                if table_text:
                    all_text.append(table_text)
            structured_text = '\n\n'.join(all_text)
        
        return {
            'text': structured_text,
            'document_structure': document_structure,
            'paragraphs': paragraphs,
            'tables': tables,
            'stats': {
                'total_elements': len(document_structure),
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables),
                'total_chars': len(structured_text),
                'has_headings': any(elem.get('is_heading', False) for elem in document_structure),
                'has_lists': any(elem.get('is_list', False) for elem in document_structure)
            }
        }
    except Exception as e:
        logger.error(f"Failed to extract DOCX content: {e}")
        raise RuntimeError(f"Failed to extract DOCX content: {str(e)}")

def _extract_table_data(table: Table) -> Dict[str, Any]:
    """Extract structured data from a DOCX table"""
    try:
        rows = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(cell for cell in row_data):  # Skip empty rows
                rows.append(row_data)
        
        if not rows:
            return None
            
        # Assume first row is header if it looks like one
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []
        
        return {
            'headers': headers,
            'rows': data_rows,
            'total_rows': len(rows),
            'total_cols': len(headers) if headers else 0
        }
    except Exception as e:
        logger.warning(f"Failed to extract table data: {e}")
        return None

def _table_to_text(table_data: Dict[str, Any]) -> str:
    """Convert table data to readable text format"""
    try:
        lines = []
        
        # Add headers
        headers = table_data.get('headers', [])
        if headers:
            lines.append(' | '.join(headers))
            lines.append('-' * 50)  # Separator
        
        # Add data rows
        for row in table_data.get('rows', []):
            lines.append(' | '.join(str(cell) for cell in row))
        
        return '\n'.join(lines)
    except Exception:
        return "[Table data could not be formatted]"

def _get_paragraph_info(para) -> Dict[str, Any]:
    """Extract detailed paragraph information"""
    try:
        style_name = para.style.name if para.style else 'Normal'
        
        return {
            'text': para.text.strip(),
            'style': style_name,
            'is_heading': style_name.startswith('Heading') or style_name.startswith('Title'),
            'heading_level': int(style_name.split()[-1]) if style_name.startswith('Heading') and style_name.split()[-1].isdigit() else None,
            'is_list': style_name.startswith('List'),
            'alignment': str(para.alignment) if para.alignment else None
        }
    except Exception:
        return {
            'text': para.text.strip() if hasattr(para, 'text') else '',
            'style': 'Normal',
            'is_heading': False,
            'heading_level': None,
            'is_list': False,
            'alignment': None
        }

def _extract_paragraph_structure(element, doc) -> Optional[Dict[str, Any]]:
    """Extract paragraph structure from XML element"""
    try:
        # Get paragraph object
        para = None
        for p in doc.paragraphs:
            if p._element == element:
                para = p
                break
        
        if para is None or not para.text.strip():
            return None
        
        para_info = _get_paragraph_info(para)
        para_info['element_type'] = 'paragraph'
        return para_info
        
    except Exception:
        return None

def _extract_table_structure(element) -> Optional[Dict[str, Any]]:
    """Extract table structure from XML element"""
    try:
        return {
            'element_type': 'table',
            'text': '[Table detected]',
            'is_table': True
        }
    except Exception:
        return None

def _generate_structured_text(document_structure: List[Dict[str, Any]]) -> str:
    """Generate text that preserves document structure"""
    try:
        lines = []
        
        for element in document_structure:
            if element['element_type'] == 'paragraph':
                if element.get('is_heading'):
                    # Add heading with appropriate formatting
                    level = element.get('heading_level', 1)
                    prefix = '#' * min(level, 6) + ' ' if level else '## '
                    lines.append(f"{prefix}{element['text']}")
                    lines.append("")  # Add blank line after heading
                elif element.get('is_list'):
                    lines.append(f"• {element['text']}")
                else:
                    lines.append(element['text'])
                    lines.append("")  # Add blank line after paragraph
            elif element['element_type'] == 'table':
                lines.append("[TABLE]")
                lines.append("")
        
        return '\n'.join(lines).strip()
    except Exception:
        return ""

def _create_structure_regions(document_structure: List[Dict[str, Any]]) -> List[Region]:
    """Create regions for document structure elements"""
    regions = []
    
    try:
        current_section = []
        current_section_type = None
        bbox_y = 0
        
        for element in document_structure:
            element_type = element.get('element_type', 'paragraph')
            
            if element.get('is_heading'):
                # Headings start new sections
                if current_section:
                    region = _create_section_region(current_section, current_section_type, bbox_y)
                    if region:
                        regions.append(region)
                        bbox_y += 20
                
                # Create heading region
                heading_region = Region(
                    region_id=str(uuid.uuid4()),
                    text=element['text'],
                    bbox=[0, bbox_y, 100, bbox_y + 10],
                    confidence=1.0,
                    metadata={
                        "engine": "python-docx",
                        "content_type": "heading",
                        "heading_level": element.get('heading_level', 1),
                        "style": element.get('style', 'Heading')
                    }
                )
                regions.append(heading_region)
                bbox_y += 10
                
                current_section = []
                current_section_type = 'content'
                
            elif element_type == 'table':
                # Tables get their own region
                if current_section:
                    region = _create_section_region(current_section, current_section_type, bbox_y)
                    if region:
                        regions.append(region)
                        bbox_y += 20
                
                current_section = []
                current_section_type = None
                
            else:
                # Regular content
                current_section.append(element)
                if current_section_type is None:
                    current_section_type = 'content'
        
        # Handle remaining content
        if current_section:
            region = _create_section_region(current_section, current_section_type, bbox_y)
            if region:
                regions.append(region)
        
    except Exception as e:
        logger.warning(f"Failed to create structure regions: {e}")
    
    return regions

def _create_section_region(section_elements: List[Dict[str, Any]], section_type: str, bbox_y: int) -> Optional[Region]:
    """Create a region for a section of content"""
    if not section_elements:
        return None
    
    try:
        section_text = '\n'.join(elem['text'] for elem in section_elements if elem.get('text'))
        
        if not section_text.strip():
            return None
        
        return Region(
            region_id=str(uuid.uuid4()),
            text=section_text,
            bbox=[0, bbox_y, 100, bbox_y + 15],
            confidence=0.8,
            metadata={
                "engine": "python-docx",
                "content_type": section_type,
                "element_count": len(section_elements),
                "has_lists": any(elem.get('is_list', False) for elem in section_elements)
            }
        )
    except Exception:
        return None
def handle_docx(doc: DocumentObject) -> BinaryDocument:
    """Handle DOCX files by extracting text and preserving document structure"""
    try:
        # Extract comprehensive content with structure
        content = extract_docx_text(doc.raw_bytes)
        
        # Create structured regions for different content types
        regions = []
        
        # Create regions for document structure elements
        if content['document_structure']:
            regions.extend(_create_structure_regions(content['document_structure']))
        
        # Create table regions
        for table_idx, table in enumerate(content['tables']):
            table_region = Region(
                region_id=str(uuid.uuid4()),
                text=_table_to_text(table),
                bbox=[0, len(regions) * 10, 100, (len(regions) + 1) * 10],
                confidence=0.9,
                metadata={
                    "engine": "python-docx",
                    "content_type": "table",
                    "table_index": table_idx,
                    "table_stats": {
                        "rows": table['total_rows'],
                        "cols": table['total_cols'],
                        "has_headers": len(table.get('headers', [])) > 0
                    }
                }
            )
            regions.append(table_region)
        
        # Fallback: create main document region if no structured regions
        if not regions and content['text']:
            regions.append(Region(
                region_id=str(uuid.uuid4()),
                text=content['text'],
                bbox=[0, 0, 100, 100],
                confidence=0.8,
                metadata={
                    "engine": "python-docx",
                    "content_type": "document_text",
                    "paragraphs": len(content['paragraphs']),
                    "tables": len(content['tables'])
                }
            ))
        
        # Wrap regions in a single Block for DOCX content
        blocks = []
        if regions:
            blocks.append(Block(
                block_id=str(uuid.uuid4()),
                title="Document Content",
                label="document",
                bbox=[0, 0, 100, 100],
                regions=regions,
                raw_text="",  # Could aggregate if needed
                confidence=0.9,
                metadata={"engine": "python-docx"}
            ))
        
        # Create page with blocks (not regions directly)
        page = Page(
            page_id=str(uuid.uuid4()),
            page_number=1,
            blocks=blocks,
            metadata={
                "docx": True,
                "extraction_stats": content['stats'],
                "has_tables": len(content['tables']) > 0,
                "has_structure": len(content['document_structure']) > 0,
                "regions_count": len(regions)
            }
        )
        
        return BinaryDocument(
            document_id=doc.document_id,
            pages=[page],
            metadata={
                "source_format": doc.detected_format,
                "mime_type": doc.mime_type,
                "extraction_method": "python-docx-structured",
                "content_stats": content['stats'],
                "file_size": len(doc.raw_bytes)
            }
        )
        
    except Exception as e:
        logger.error(f"Failed to handle DOCX document {doc.document_id}: {e}")
        # Return empty document with error info
        return BinaryDocument(
            document_id=doc.document_id,
            pages=[],
            metadata={
                "source_format": doc.detected_format,
                "mime_type": doc.mime_type,
                "error": str(e),
                "file_size": len(doc.raw_bytes)
            }
        )